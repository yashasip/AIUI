import docx
import PyPDF2
import requests


class TextDataFileHandle:
    '''Handles Pdf, Docx and Direct link to pdfs
    type: '''
    def __init__(self, filePath) -> None:
        self.filePath = filePath
        self.fileType = self.getFileType()

        self.setupTextDataFile()

    def getFileType(self):
        for index, character in enumerate(self.filePath[::-1]):
            if character == ".":
                return self.filePath[-index:].lower()

    def setupTextDataFile(self): # setups the text data file
        if self.filePath.startswith("https://") or self.filePath.startswith("http://"):
            self.scrapePdf() # scrape pdf when link is given
            self.filePath = "inputPDF.pdf"
        if self.fileType == "pdf": # also called after scraping and when pdf file is chosen 
            pdfFile = open(self.filePath, "rb")
            self.reader = PyPDF2.PdfFileReader(pdfFile)
            self.pageCount = self.reader.numPages

    def readTextDataFile(self): # reads text data file based on file type
        if self.fileType == "pdf":
            textData = self.readPDF()
        elif self.fileType == "docx":
            textData = self.readDocx()
        return textData

    def readPDF(self, start=0): # extracts pdf data, returns it as string
        pdfData = ""
        for i in range(start, self.pageCount):
            page = self.reader.getPage(i)
            pdfData = pdfData + str(page.extractText())

        return pdfData

    def readDocx(self): # extracts document data, returns it as a string
        doc = docx.Document(self.filePath)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return "\n".join(fullText)

    def scrapePdf(self): # scrapes pdf from direct link
        html = requests.get(self.filePath, stream=True)
        file = open("inputPDF.pdf", "wb") # create a new pdf
        for chunk in html.iter_content(1000):
            file.write(chunk)
        file.close()

    def saveTextDataFile(self,path, content): # saves as doc always
        newDoc = docx.Document() # creates new document
        newDoc.add_paragraph(content)
        newDoc.save(path)        